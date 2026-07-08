"""
reporting/report_generator.py
-------------------------------
Assembles the Daily Trading Report Word document from processed DataFrames
and converts it to PDF.
"""

from __future__ import annotations

import subprocess
import platform
from datetime import date
from pathlib import Path
from typing import NamedTuple

import pandas as pd
from docx import Document

from config import (
    B_PREFIX,
    DATE_COL,
    OUTPUT_DIR,
    WORD_TEMPLATE_PATH,
)
from reporting.word_styles import (
    add_normal,
    add_subtitle,
    add_title,
    set_cell_green,
    set_header_row_style,
)
from utils.logger import get_logger

logger = get_logger(__name__)


# ---------------------------------------------------------------------------
# Data container
# ---------------------------------------------------------------------------


class ReportData(NamedTuple):
    """All processed data needed to render the report."""

    movers: dict[str, pd.DataFrame]
    """Strategy name → filtered Recent-Movers DataFrame (only non-empty sheets)."""

    recent_sectors: pd.DataFrame
    """Date, Ticker, Category/Industry, Subcategory, RR."""

    tendencies_ranking: pd.DataFrame
    """Indicator, Count — green based on latest-day value in source."""

    tendencies_ranking_source: pd.DataFrame
    """Full Tendencies DataFrame used to determine latest-day green highlight."""

    tendencies_c: pd.DataFrame
    """Date + C_ columns — no green highlighting."""


# ---------------------------------------------------------------------------
# Public entry-point
# ---------------------------------------------------------------------------


def generate_report(data: ReportData) -> Path:
    """Build the Word document, save it, convert to PDF, return PDF path."""
    doc = _open_template()

    add_title(doc, "Daily Trading Report")
    add_normal(doc, f"Generated on {date.today().isoformat()}")

    # Section 1
    add_subtitle(doc, "Recent Movers")
    _write_movers_section(doc, data.movers)

    # Section 2
    add_subtitle(doc, "Recent Sectors")
    _write_dataframe_table(doc, data.recent_sectors, highlight_mode="none")

    # Section 3/4 – Ranking with latest-day green
    add_subtitle(doc, "Tendencies")
    _write_ranking_table(doc, data.tendencies_ranking, data.tendencies_ranking_source)

    # Tendencies Detail – no green
    add_subtitle(doc, "Tendencies Detail")
    _write_dataframe_table(doc, data.tendencies_c, highlight_mode="none")

    docx_path = _save_document(doc)
    pdf_path  = _convert_to_pdf(docx_path)

    return pdf_path if pdf_path else docx_path


# ---------------------------------------------------------------------------
# Section writers
# ---------------------------------------------------------------------------


def _write_movers_section(doc: Document, movers: dict[str, pd.DataFrame]) -> None:
    """Write one sub-section per strategy sheet.
    Sheets with no data are already excluded before this point."""
    if not movers:
        add_normal(doc, "No data available for the selected period.")
        return

    for sheet_name, df in movers.items():
        # df is guaranteed non-empty (empty sheets were skipped in main.py)
        add_normal(doc, f"{sheet_name} Table")
        _write_dataframe_table(doc, df, highlight_mode="none")


def _write_ranking_table(
    doc: Document,
    ranking: pd.DataFrame,
    source_df: pd.DataFrame,
) -> None:
    """Render the Tendencies Ranking table.

    Green highlight rule (change #6):
    - Find the most-recent date in *source_df*.
    - For each indicator, check its value on that date in *source_df*.
    - If the value is > 0 → green; otherwise → no highlight.
    - Counts are NOT modified.
    """
    if ranking is None or ranking.empty:
        add_normal(doc, "  (no data)")
        return

    # Build {indicator_name: value_on_latest_day} lookup
    latest_day_values: dict[str, float] = {}
    if not source_df.empty and DATE_COL in source_df.columns:
        latest_date = source_df[DATE_COL].max()
        latest_row = source_df.loc[source_df[DATE_COL] == latest_date]
        if not latest_row.empty:
            for col in source_df.columns:
                if col.startswith(B_PREFIX):
                    val = latest_row.iloc[0][col]
                    try:
                        latest_day_values[col] = float(val)
                    except (TypeError, ValueError):
                        latest_day_values[col] = 0.0

    n_rows, n_cols = ranking.shape
    table = doc.add_table(rows=1 + n_rows, cols=n_cols)
    table.style = "Table Grid"

    # Header
    for col_idx, col_name in enumerate(ranking.columns):
        cell = table.rows[0].cells[col_idx]
        cell.text = str(col_name)
        set_header_row_style(cell)

    # Data rows
    count_col_idx = list(ranking.columns).index("Count") if "Count" in ranking.columns else -1

    for row_idx, (_, series) in enumerate(ranking.iterrows()):
        word_row = table.rows[row_idx + 1].cells
        indicator = str(series.get("Indicator", ""))
        latest_val = latest_day_values.get(indicator, 0.0)
        apply_green = latest_val > 0

        for col_idx, col_name in enumerate(ranking.columns):
            raw = series[col_name]
            cell = word_row[col_idx]
            cell.text = _format_cell_value(raw)
            if apply_green:
                set_cell_green(cell)

    doc.add_paragraph("")


# ---------------------------------------------------------------------------
# Generic DataFrame → Word table
# ---------------------------------------------------------------------------


def _write_dataframe_table(
    doc: Document,
    df: pd.DataFrame,
    highlight_mode: str = "none",
) -> None:
    """Render *df* as a Word table.

    highlight_mode:
        ``"latest_day"`` – green on numeric > 0 cells in the most-recent date row.
        ``"none"``        – no highlighting.
    """
    if df is None or df.empty:
        add_normal(doc, "  (no data)")
        return

    n_rows, n_cols = df.shape
    table = doc.add_table(rows=1 + n_rows, cols=n_cols)
    table.style = "Table Grid"

    latest_date: pd.Timestamp | None = None
    if highlight_mode == "latest_day" and DATE_COL in df.columns:
        latest_date = pd.to_datetime(df[DATE_COL]).max()

    # Header
    for col_idx, col_name in enumerate(df.columns):
        cell = table.rows[0].cells[col_idx]
        cell.text = str(col_name)
        set_header_row_style(cell)

    # Data
    for row_idx, (_, series) in enumerate(df.iterrows()):
        word_row = table.rows[row_idx + 1].cells
        is_latest = (
            latest_date is not None
            and DATE_COL in df.columns
            and pd.to_datetime(series[DATE_COL]) == latest_date
        )
        for col_idx, col_name in enumerate(df.columns):
            raw = series[col_name]
            cell = word_row[col_idx]
            cell.text = _format_cell_value(raw)
            if is_latest and _is_positive(raw):
                set_cell_green(cell)

    doc.add_paragraph("")


# ---------------------------------------------------------------------------
# PDF conversion
# ---------------------------------------------------------------------------


def _convert_to_pdf(docx_path: Path) -> Path | None:
    """Convert *docx_path* to PDF using LibreOffice (cross-platform).

    On Windows, falls back to the Microsoft Word COM interface when
    LibreOffice is not found.

    Returns the PDF path on success, or *None* on failure.
    """
    pdf_path = docx_path.with_suffix(".pdf")

    # --- Try LibreOffice first (works on Windows, macOS, Linux) ---
    libreoffice_candidates = [
        "soffice",                                          # Linux / macOS PATH
        r"C:\Program Files\LibreOffice\program\soffice.exe",  # Windows default
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for soffice in libreoffice_candidates:
        try:
            result = subprocess.run(
                [
                    soffice,
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", str(docx_path.parent),
                    str(docx_path),
                ],
                capture_output=True,
                text=True,
                timeout=60,
            )
            if result.returncode == 0 and pdf_path.exists():
                logger.info("PDF created via LibreOffice → %s", pdf_path)
                return pdf_path
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue

    # --- Fallback: Microsoft Word COM (Windows only) ---
    if platform.system() == "Windows":
        try:
            import comtypes.client  # type: ignore

            word = comtypes.client.CreateObject("Word.Application")
            word.Visible = False
            doc_com = word.Documents.Open(str(docx_path.resolve()))
            doc_com.SaveAs(str(pdf_path.resolve()), FileFormat=17)  # 17 = wdFormatPDF
            doc_com.Close()
            word.Quit()
            if pdf_path.exists():
                logger.info("PDF created via Microsoft Word COM → %s", pdf_path)
                return pdf_path
        except Exception as exc:
            logger.warning("Word COM conversion failed: %s", exc)

    logger.warning(
        "PDF conversion unavailable. "
        "Install LibreOffice (https://www.libreoffice.org) or ensure Microsoft Word is present. "
        "The .docx report is still available at %s",
        docx_path,
    )
    return None


# ---------------------------------------------------------------------------
# Private helpers
# ---------------------------------------------------------------------------


def _open_template() -> Document:
    if WORD_TEMPLATE_PATH.exists():
        logger.info("Using Word template: %s", WORD_TEMPLATE_PATH)
        return Document(str(WORD_TEMPLATE_PATH))
    logger.warning(
        "Template not found at '%s'; generating without custom styles.",
        WORD_TEMPLATE_PATH,
    )
    return Document()


def _save_document(doc: Document) -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    filename = f"Daily_Report_{date.today().isoformat()}.docx"
    output_path = OUTPUT_DIR / filename
    doc.save(str(output_path))
    logger.info("Word document saved → %s", output_path)
    return output_path


def _format_cell_value(value: object) -> str:
    if pd.isna(value):
        return ""
    if isinstance(value, pd.Timestamp):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, float) and value == int(value):
        return str(int(value))
    return str(value)


def _is_positive(value: object) -> bool:
    try:
        return float(value) > 0  # type: ignore[arg-type]
    except (TypeError, ValueError):
        return False
