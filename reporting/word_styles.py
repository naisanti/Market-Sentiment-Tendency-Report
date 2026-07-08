"""
reporting/word_styles.py
-------------------------
Low-level python-docx helpers for applying styles, colours, and cell
formatting.  All direct python-docx API calls live here; the report
generator only calls these helpers.
"""

from __future__ import annotations

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.table import _Cell

from config import GREEN_FILL, STYLE_NORMAL, STYLE_SUBTITLE, STYLE_TITLE
from utils.logger import get_logger

logger = get_logger(__name__)


# ---------------------------------------------------------------------------
# Document-level helpers
# ---------------------------------------------------------------------------


def add_title(doc: Document, text: str) -> None:
    """Add a paragraph using the Title Rock style."""
    _safe_add_paragraph(doc, text, STYLE_TITLE)


def add_subtitle(doc: Document, text: str) -> None:
    """Add a paragraph using the Subtitle Rock style."""
    _safe_add_paragraph(doc, text, STYLE_SUBTITLE)


def add_normal(doc: Document, text: str) -> None:
    """Add a paragraph using the Normal Trading Inform style."""
    _safe_add_paragraph(doc, text, STYLE_NORMAL)


# ---------------------------------------------------------------------------
# Cell-level helpers
# ---------------------------------------------------------------------------


def set_cell_green(cell: _Cell) -> None:
    """Apply a light-green background fill to *cell*."""
    _set_cell_background(cell, GREEN_FILL)


def set_cell_background(cell: _Cell, hex_color: str) -> None:
    """Apply an arbitrary background fill (hex, no '#') to *cell*."""
    _set_cell_background(cell, hex_color)


def set_cell_bold(cell: _Cell) -> None:
    """Make all runs in *cell* bold."""
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True


def set_header_row_style(
    cell: _Cell,
    bg_hex: str = "2E4057",
    font_color: tuple[int, int, int] = (255, 255, 255),
) -> None:
    """Dark background + white bold text — standard header row style."""
    _set_cell_background(cell, bg_hex)
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(*font_color)


# ---------------------------------------------------------------------------
# Private helpers
# ---------------------------------------------------------------------------


def _safe_add_paragraph(doc: Document, text: str, style_name: str) -> None:
    """Add a paragraph with *style_name*, falling back to Normal if missing."""
    try:
        doc.add_paragraph(text, style=style_name)
    except KeyError:
        logger.warning(
            "Style '%s' not found in the template; using built-in 'Normal'.",
            style_name,
        )
        doc.add_paragraph(text)


def _set_cell_background(cell: _Cell, hex_color: str) -> None:
    """Set the <w:shd> element on a table cell to a solid fill colour."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    # Remove any existing shading element to avoid duplicates
    existing = tc_pr.find(qn("w:shd"))
    if existing is not None:
        tc_pr.remove(existing)

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.upper())
    tc_pr.append(shd)
